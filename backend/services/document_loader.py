from pathlib import Path
from pypdf import PdfReader
from docx import Document
import csv


def load_pdf(file_path: str):
    pages = []

    pdf = PdfReader(file_path)

    for page_num, page in enumerate(pdf.pages, start=1):
        text = page.extract_text() or ""

        pages.append({
            "text": text,
            "page": page_num
        })

    return pages


def load_docx(file_path: str):
    """
    Robust DOCX loader.

    Extracts text from:
    - Paragraphs
    - Tables

    This improves extraction for resumes and structured
    Word documents where important information may be
    stored inside tables.
    """

    doc = Document(file_path)

    parts = []

    # -----------------------------
    # Extract normal paragraphs
    # -----------------------------
    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # -----------------------------
    # Extract tables
    # -----------------------------
    for table in doc.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    row_text.append(text)

            if row_text:
                parts.append(" | ".join(row_text))

    return [{
        "text": "\n".join(parts),
        "page": 1
    }]


def load_txt(file_path: str):

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    return [{
        "text": text,
        "page": 1
    }]


def load_csv(file_path: str):

    rows = []

    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:

        reader = csv.reader(f)

        for row in reader:
            rows.append(" | ".join(row))

    return [{
        "text": "\n".join(rows),
        "page": 1
    }]


def load_document(file_path: str):

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return load_pdf(file_path)

    if ext == ".docx":
        return load_docx(file_path)

    if ext in [".txt", ".md"]:
        return load_txt(file_path)

    if ext == ".csv":
        return load_csv(file_path)

    raise ValueError(f"Unsupported file type: {ext}")

